"""
Builds ex_Diary_melnik.docx from the Сладков template.
Replaces: name, dates, calendar plan table, characteristic text,
          individual task body, supervisor review body.
"""
from docx import Document

SRC  = "f:/Adam-Chip/diploma/Practicum/ex_Diary.docx"
DEST = "f:/Adam-Chip/diploma/Practicum/ex_Diary_melnik.docx"


# ── helpers ────────────────────────────────────────────────────────────────

def para_text(p):
    return "".join(r.text for r in p.runs)

def replace_in_para(para, old, new):
    full = para_text(para)
    if old not in full:
        return False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    new_full = full.replace(old, new)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    return True

def replace_all_paras(doc, subs):
    for para in doc.paragraphs:
        for old, new in subs:
            replace_in_para(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in subs:
                        replace_in_para(para, old, new)

def set_para_text(para, text):
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)

def set_cell_text(cell, text):
    """Clear cell and write plain text preserving first-run formatting."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if cell.paragraphs:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].text = text
        else:
            cell.paragraphs[0].add_run(text)
    else:
        cell.add_paragraph(text)


# ── content ────────────────────────────────────────────────────────────────

SUBS = [
    ("Сладков Даниила Павловича",       "Мельника Романа Игоревича"),
    ("Сладков Даниил Павлович",         "Мельник Роман Игоревич"),
    ("Сладкова Даниила Павловича",      "Мельника Романа Игоревича"),
    ("Сладков Д.П.",                    "Мельник Р.И."),
    ("Сладков Д. П.",                   "Мельник Р.И."),
    ("Сладков",                         "Мельник"),
    ("Даниил Павлович",                 "Роман Игоревич"),
    ("Даниила Павловича",               "Романа Игоревича"),
    ("28» апреля 2024г.",               "27» апреля 2025г."),
    ("28» сентября 2025г.",             "27» апреля 2025г."),
    ("25» мая 2025г.",                  "24» мая 2025г."),
    ("с «28» апреля 2024г. по «25» мая 2025г.",
     "с «27» апреля 2025г. по «24» мая 2025г."),
    ("28.04.2025 – 30.04.2025",         "27.04.2025"),
    (".05.2025 – 25.05.2025",           "24.05.2025"),
]

CALENDAR = [
    ("27.04.2025",
     "Получение индивидуального задания."),
    ("28.04.2025 – 29.04.2025",
     "Анализ философской и технической литературы по теме цифровой субъектности "
     "и агентных архитектур (Хейлз, Чалмерс, Деннет, Гофман, Дарем Питерс)."),
    ("30.04.2025",
     "Разработка восьмикритериальной матрицы квазисубъектности. "
     "Сравнительный анализ агентных систем (Agent Ruby, Being, Voyager, "
     "BOB, Generative Agents, Symbiosis of Agents)."),
    ("01.05.2025 – 09.05.2025",
     "Проектирование модульной архитектуры агента. Реализация инференс-стека: "
     "LLM (Gemma 4 E4B), VLM (VILA 1.5), ASR (WhisperX), TTS (Silero) "
     "на NVIDIA Jetson Orin NX. Разработка оркестратора (FastAPI + asyncio)."),
    ("12.05.2025 – 16.05.2025",
     "Разработка системы памяти, когнитивного цикла, командного контура. "
     "Интеграция ESP32-S3. Речевой цикл (OpenWakeWord, WebRTC VAD). "
     "Веб-интерфейс оператора."),
    ("19.05.2025 – 20.05.2025",
     "3D-моделирование элементов инсталляции: генерация референсов "
     "(gpt-image-2, Tripo 3D), скульптинг-коррекция дефектов генерации в Blender."),
    ("21.05.2025",
     "Параметрическое и ручное моделирование: шейные позвонки, корпуса "
     "вибромоторов, динамиков, микрофонов. Подготовка файлов и 3D-печать "
     "элементов инсталляции."),
    ("22.05.2025 – 23.05.2025",
     "Стендовое тестирование системы (98 сессий, 902 цикла). "
     "Сбор и систематизация метрик поведенческой согласованности агента."),
    ("24.05.2025",
     "Систематизация результатов практики. Написание отчёта. Заполнение дневника."),
]

TASK_NEW = (
    "В рамках преддипломной практики Мельнику Роману Игоревичу ставится задача "
    "по комплексной работе над дипломным проектом — нейроагентной художественной "
    "инсталляцией «Адам Чип». Задание включает: анализ теоретической базы по теме "
    "цифровой субъектности и агентных архитектур; разработку восьмикритериальной "
    "аналитической матрицы квазисубъектности и сравнительный анализ шести агентных "
    "систем (Agent Ruby, Being, Voyager, BOB, Generative Agents, Symbiosis of Agents); "
    "проектирование и программную реализацию модульной архитектуры агента "
    "(языковое ядро LLM, эпизодическая память, когнитивный цикл, сенсорный и "
    "моторный слои); 3D-моделирование, скульптинг и 3D-печать элементов физической "
    "инсталляции; стендовое тестирование системы и сбор метрик поведенческой "
    "согласованности агента."
)

REVIEW_NEW = (
    "Мельник Роман Игоревич успешно прошёл преддипломную практику, "
    "продемонстрировав высокий уровень профессиональной подготовки, "
    "самостоятельность и выраженный аналитический подход к решению поставленных задач. "
    "В ходе практики Роман Игоревич провёл значительную работу по разработке и "
    "программной реализации нейроагентной системы «Адам Чип», что потребовало "
    "глубоких знаний в области программирования, искусственного интеллекта и "
    "встраиваемых систем. Была выполнена разработка восьмикритериальной аналитической "
    "матрицы квазисубъектности и сравнительный анализ агентных систем. "
    "Кроме того, Роман Игоревич качественно выполнил задачу по 3D-моделированию и "
    "3D-печати элементов физической инсталляции, показав понимание принципов "
    "конструктивного проектирования и работы с генеративными инструментами. "
    "Мельник Р.И. проявил себя как ответственный, инициативный и компетентный "
    "практикант, способный эффективно организовывать работу над сложным "
    "техническим и творческим проектом."
)


# ── open ──────────────────────────────────────────────────────────────────

doc = Document(SRC)

# ── Phase 1: simple text substitutions everywhere ─────────────────────────

replace_all_paras(doc, SUBS)

# ── Phase 2: calendar plan table ──────────────────────────────────────────
# Nested table: doc.tables[1].rows[1].cells[0].tables[0]

try:
    cal_table = doc.tables[1].rows[1].cells[0].tables[0]
    print(f"Found calendar nested table: {len(cal_table.rows)} rows")

    data_rows = cal_table.rows[1:]  # skip header

    for i, (date_str, work_str) in enumerate(CALENDAR):
        if i < len(data_rows):
            row = data_rows[i]
            cells = row.cells
            set_cell_text(cells[0], date_str)
            set_cell_text(cells[1], work_str)
            if len(cells) > 2:
                set_cell_text(cells[2], "")
        else:
            print(f"  Warning: no row for calendar entry {i}: {date_str}")

    for i in range(len(CALENDAR), len(data_rows)):
        for cell in data_rows[i].cells:
            set_cell_text(cell, "")

    print("Calendar updated.")

except Exception as e:
    print(f"WARNING: calendar table update failed: {e}")


# ── Phase 3: replace Характеристика body ─────────────────────────────────
# doc.tables[1].rows[0].cells[2]
# Structure: [blank][header][blank][name line][(фамилия, имя, отчество)][body paras...]

def replace_characteristic_body(cell, review_new):
    """Replace characteristic body paragraphs; preserve header, name line, label."""
    paras = cell.paragraphs
    body_start = None
    for i, para in enumerate(paras):
        if "фамилия, имя, отчество" in para_text(para):
            body_start = i + 1
            break
    if body_start is None:
        print("WARNING: Характеристика body anchor not found")
        return False
    body_indices = [i for i in range(body_start, len(paras))
                    if para_text(paras[i]).strip()]
    if not body_indices:
        print("WARNING: no body paragraphs found in Характеристика cell")
        return False
    set_para_text(paras[body_indices[0]], review_new)
    for i in body_indices[1:]:
        set_para_text(paras[i], "")
    print(f"Характеристика updated: {len(body_indices)} paragraphs replaced")
    return True

try:
    replace_characteristic_body(doc.tables[1].rows[0].cells[2], REVIEW_NEW)
except Exception as e:
    print(f"WARNING: Characteristic cell update failed: {e}")


# ── Phase 4: replace task body + review body ──────────────────────────────
# doc.tables[2].rows[0].cells[1]
# Structure: [blank][ИНДИВИДУАЛЬНОЕ ЗАДАНИЕ header][blank][task para]
#            [blank][Отзыв руководителя header][blank]
#            [review paras...][last review para + "Оценка выполнения..."]
#            [blank][blank][signature lines...]

def replace_task_and_review(cell, task_new, review_new):
    """
    Replace task body and review body; preserve all headers and signature lines.
    Handles two layouts:
      A) Review in separate paragraphs, Оценка mixed into the last one.
      B) Entire review + Оценка in a single paragraph (line-break separated runs).
    """
    from docx.oxml.ns import qn as _qn

    paras = cell.paragraphs
    task_idx = None
    after_review_header = False
    review_body = []
    ozenka_idx = None

    for i, para in enumerate(paras):
        text = para_text(para)
        stripped = text.strip()

        if task_idx is None and "В рамках" in text and "задача" in text:
            task_idx = i
        elif "Отзыв руководителя" in text:
            after_review_header = True
        elif after_review_header and ozenka_idx is None:
            if "Оценка выполнения" in text:
                ozenka_idx = i
            elif (stripped
                  and "___" not in stripped
                  and "(подпись)" not in text
                  and "(фамилия" not in text
                  and "Руководитель" not in text):
                review_body.append(i)

    if task_idx is not None:
        set_para_text(paras[task_idx], task_new)
        print(f"Task body replaced at para {task_idx}")
    else:
        print("WARNING: task body paragraph not found")

    if review_body:
        # Layout A: review in separate paragraphs
        set_para_text(paras[review_body[0]], review_new)
        for i in review_body[1:]:
            set_para_text(paras[i], "")
        print(f"Review body replaced: {len(review_body)} paragraphs")
        if ozenka_idx is not None:
            set_para_text(paras[ozenka_idx],
                          "Оценка выполнения индивидуального задания ______________________________")
            print(f"Оценка paragraph fixed at para {ozenka_idx}")

    elif ozenka_idx is not None:
        # Layout B: review text and Оценка are in ONE paragraph with w:br separators
        ozenka_para = paras[ozenka_idx]
        ozenka_run_start = None
        for i, run in enumerate(ozenka_para.runs):
            if "Оценка выполнения" in run.text:
                ozenka_run_start = i
                break

        if ozenka_run_start is not None and ozenka_run_start > 0:
            # Clear runs before the Оценка run; strip w:br elements too
            for run in ozenka_para.runs[:ozenka_run_start]:
                for br in run._r.findall(_qn("w:br")):
                    run._r.remove(br)
                run.text = ""
            ozenka_para.runs[0].text = review_new
            print(f"Combined review+Оценка fixed at para {ozenka_idx}, "
                  f"Оценка at run {ozenka_run_start}")
        else:
            # Phase 1 collapsed all runs into run 0 — remove stale w:br elements
            for run in ozenka_para.runs:
                for br in run._r.findall(_qn("w:br")):
                    run._r.remove(br)
            set_para_text(ozenka_para,
                          review_new
                          + " Оценка выполнения индивидуального задания ______________________________")
            print(f"Combined review+Оценка (fallback) at para {ozenka_idx}")
    else:
        print("WARNING: review body paragraphs not found")

try:
    replace_task_and_review(doc.tables[2].rows[0].cells[1], TASK_NEW, REVIEW_NEW)
except Exception as e:
    print(f"WARNING: Task/review cell update failed: {e}")


# ── save ──────────────────────────────────────────────────────────────────

doc.save(DEST)
print(f"Saved to: {DEST}")
