#!/usr/bin/env python3
"""
Post-process Pandoc-generated diploma.docx:
- Tables: autofit layout, standard 0.5pt borders, no cell paragraph indent
- Captions: remove italic/bold from all runs
"""
import sys
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT = os.path.join(SCRIPT_DIR, "diploma.docx")


def _set_tbl_autofit(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Remove existing layout element and replace with autofit
    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit")
    tblPr.append(layout)

    # Set table width to 100% of text body (prevents overflow beyond page margins).
    # w:type="pct" uses units of 1/50 of a percent, so 5000 = 100%.
    for tblW in tblPr.findall(qn("w:tblW")):
        tblPr.remove(tblW)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)

    # Remove fixed column widths from tblGrid
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        for col in tblGrid.findall(qn("w:gridCol")):
            col.attrib.pop(qn("w:w"), None)

    # Set table-level borders
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")       # 0.5pt
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _fix_cell(cell):
    tc = cell._tc

    # Remove fixed cell width (let autofit decide)
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is not None:
        for tcW in tcPr.findall(qn("w:tcW")):
            tcW.set(qn("w:type"), "auto")
            tcW.set(qn("w:w"), "0")

    # Remove first-line indent from each paragraph in cell
    for para in cell.paragraphs:
        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None:
            for ind in pPr.findall(qn("w:ind")):
                ind.attrib.pop(qn("w:firstLine"), None)
                ind.attrib.pop(qn("w:firstLineChars"), None)
        para.paragraph_format.first_line_indent = Cm(0)


def _fix_header_row_border(table):
    """Explicitly set 1pt bottom border on header row cells (overrides tblHeader quirk)."""
    if not table.rows:
        return
    first_row = table.rows[0]
    for cell in first_row.cells:
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        for old in tcBorders.findall(qn("w:bottom")):
            tcBorders.remove(old)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")    # 1pt — visually separates header from data rows
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), "000000")
        tcBorders.append(bottom)


def process_tables(doc):
    count = 0
    for table in doc.tables:
        _set_tbl_autofit(table)
        _fix_header_row_border(table)
        for row in table.rows:
            for cell in row.cells:
                _fix_cell(cell)
        count += 1
    return count


def fix_captions(doc):
    count = 0
    for para in doc.paragraphs:
        if para.style.name == "Caption":
            para.paragraph_format.first_line_indent = Cm(0)
            for run in para.runs:
                run.font.italic = False
                run.font.bold = False
            count += 1
    return count


def fix_table_titles(doc):
    """Remove first-line indent from paragraphs immediately preceding tables."""
    count = 0
    body = doc.element.body
    elems = list(body)
    for i, el in enumerate(elems):
        if el.tag.split("}")[-1] == "tbl" and i > 0:
            prev = elems[i - 1]
            if prev.tag.split("}")[-1] == "p":
                pPr = prev.find(qn("w:pPr"))
                if pPr is None:
                    pPr = OxmlElement("w:pPr")
                    prev.insert(0, pPr)
                for old in pPr.findall(qn("w:ind")):
                    pPr.remove(old)
                ind = OxmlElement("w:ind")
                ind.set(qn("w:firstLine"), "0")
                pPr.append(ind)
                count += 1
    return count


def _apply_tnr_font(rPr):
    """Set Times New Roman in all font slots; strip theme-font overrides. Does not touch size."""
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    fonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts.set(qn(attr), "Times New Roman")
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
        fonts.attrib.pop(qn(attr), None)
    rPr.insert(0, fonts)


def _set_tnr_14(rPr):
    """Overwrite font and size in an rPr element: Times New Roman 14pt, all slots."""
    _apply_tnr_font(rPr)
    for old in rPr.findall(qn("w:sz")):
        rPr.remove(old)
    for old in rPr.findall(qn("w:szCs")):
        rPr.remove(old)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "28")
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "28")
    rPr.append(sz)
    rPr.append(szCs)


def fix_heading_fonts(doc):
    """Force Times New Roman font on all Heading styles (keeps style-defined sizes)."""
    fixed = 0
    for style in doc.styles:
        if style.name.startswith("Heading"):
            rPr = style.element.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                style.element.append(rPr)
            _apply_tnr_font(rPr)
            fixed += 1
    return fixed


def fix_verbatim_font_size(doc):
    """Force Times New Roman 14pt on all code styles (Verbatim Char, Source Code)."""
    fixed = 0
    for style in doc.styles:
        if style.name in ("Verbatim Char", "Source Code"):
            rPr = style.element.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                style.element.append(rPr)
            _set_tnr_14(rPr)
            fixed += 1
    return fixed


def remove_bookmarks(doc):
    count = 0
    body = doc.element.body
    for tag in (qn("w:bookmarkStart"), qn("w:bookmarkEnd")):
        for bm in body.findall(".//" + tag):
            parent = bm.getparent()
            if parent is not None:
                parent.remove(bm)
                count += 1
    return count


def add_appendix_page_breaks(doc):
    """Insert page-break-before on every heading that starts with «Приложение»."""
    count = 0
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading") and para.text.strip().startswith("Приложение"):
            pPr = para._p.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                para._p.insert(0, pPr)
            for old in pPr.findall(qn("w:pageBreakBefore")):
                pPr.remove(old)
            pb = OxmlElement("w:pageBreakBefore")
            pb.set(qn("w:val"), "true")
            pPr.append(pb)
            count += 1
    return count


def main():
    if not os.path.exists(INPUT):
        print(f"FAIL: {INPUT} not found")
        sys.exit(1)

    doc = Document(INPUT)
    tables = process_tables(doc)
    captions = fix_captions(doc)
    titles = fix_table_titles(doc)
    fix_verbatim_font_size(doc)
    headings = fix_heading_fonts(doc)
    bookmarks = remove_bookmarks(doc)
    appendices = add_appendix_page_breaks(doc)
    doc.save(INPUT)
    print(f"OK: postprocessed {INPUT} ({tables} tables, {captions} captions, {titles} table titles, {headings} heading styles fixed, {bookmarks} bookmarks removed, {appendices} appendix breaks)")


if __name__ == "__main__":
    main()
